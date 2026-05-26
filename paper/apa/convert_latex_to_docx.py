#!/usr/bin/env python3
import os
import re
import glob
import subprocess
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths configuration
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MAIN_TEX = os.path.join(BASE_DIR, "thesis_apa.tex")
PREPROCESSED_TEX = os.path.join(BASE_DIR, "thesis_preprocessed.tex")
DRAFT_DOCX = os.path.join(BASE_DIR, "thesis_draft.docx")
FINAL_DOCX = os.path.join(BASE_DIR, "thesis_apa.docx")
BIB_FILE = os.path.join(BASE_DIR, "references.bib")
CSL_FILE = os.path.join(BASE_DIR, "apa.csl")
ICREATE_PDF = os.path.join(BASE_DIR, "iCreate.pdf")

# Monospace detection sets (used in is_mono_run)
_MONO_FONT_KEYWORDS = {
    'courier', 'consolas', 'lucida console', 'andale mono',
    'source code pro', 'dejavu sans mono', 'monaco', 'menlo',
}
_MONO_STYLE_KEYWORDS = {'verbatim', 'code', 'inline code', 'source code'}

# Math namespace
_MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def convert_pdf_to_png(pdf_path, png_path, dpi=300):
    """Converts a single-page PDF to PNG using pdftoppm."""
    base = os.path.splitext(png_path)[0]
    cmd = ["pdftoppm", "-png", "-r", str(dpi), "-singlefile", pdf_path, base]
    try:
        subprocess.run(cmd, check=True)
        print(f"Successfully converted: {pdf_path} -> {png_path}")
    except subprocess.CalledProcessError as e:
        print(f"Error converting {pdf_path}: {e}")


def convert_multipage_pdf_to_pngs(pdf_path, output_prefix, dpi=200):
    """Converts a multi-page PDF to per-page PNGs (output_prefix-1.png, etc.)."""
    cmd = ["pdftoppm", "-png", "-r", str(dpi), pdf_path, output_prefix]
    try:
        subprocess.run(cmd, check=True)
        print(f"Successfully split multipage PDF: {pdf_path} into pages with prefix {output_prefix}")
    except subprocess.CalledProcessError as e:
        print(f"Error splitting multipage PDF {pdf_path}: {e}")


def inline_latex(filepath, base_dir):
    """Recursively inlines all \\input{...} statements in a LaTeX file."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    def replacer(match):
        filename = match.group(1).strip()
        filename = filename.replace('"', '').replace("'", "")
        if not filename.endswith('.tex'):
            input_path = os.path.join(base_dir, filename + '.tex')
            if not os.path.exists(input_path):
                input_path = os.path.join(base_dir, filename)
        else:
            input_path = os.path.join(base_dir, filename)

        if os.path.exists(input_path):
            print(f"Inlining: {input_path}")
            return inline_latex(input_path, base_dir)
        else:
            print(f"Warning: input file not found: {input_path}")
            return match.group(0)

    pattern = re.compile(r'\\input\s*\{([^}]+)\}')
    while pattern.search(content):
        content = pattern.sub(replacer, content)
    return content


def strip_resizebox(text):
    """Removes all \\resizebox{\\width}{!}{...} wrappers, keeping inner content."""
    idx = 0
    while True:
        idx = text.find(r'\resizebox', idx)
        if idx == -1:
            break

        brace_count = 0
        brace_start_idx = -1
        search_idx = idx + len(r'\resizebox')

        while search_idx < len(text):
            if text[search_idx] == '{':
                brace_count += 1
                if brace_count == 3:
                    brace_start_idx = search_idx
                    break
            search_idx += 1

        if brace_start_idx == -1:
            idx += len(r'\resizebox')
            continue

        stack = 1
        match_end_idx = -1
        for i in range(brace_start_idx + 1, len(text)):
            if text[i] == '{':
                stack += 1
            elif text[i] == '}':
                stack -= 1
                if stack == 0:
                    match_end_idx = i
                    break

        if match_end_idx == -1:
            idx += len(r'\resizebox')
            continue

        inner_content = text[brace_start_idx + 1:match_end_idx]
        if inner_content.startswith('%'):
            inner_content = inner_content[1:]
        inner_content = inner_content.strip()

        text = text[:idx] + inner_content + text[match_end_idx + 1:]

    return text


def _trim_png_whitespace(png_path, border_px=15):
    """Uses PIL to auto-crop white margins from a PNG, adding a small border."""
    from PIL import Image, ImageOps
    try:
        img = Image.open(png_path).convert('RGB')
        inverted = ImageOps.invert(img.convert('L'))
        bbox = inverted.getbbox()
        if bbox:
            x1, y1, x2, y2 = bbox
            w, h = img.size
            img = img.crop((max(0, x1 - border_px), max(0, y1 - border_px),
                            min(w, x2 + border_px), min(h, y2 + border_px)))
        img.save(png_path)
        print(f"Trimmed whitespace: {png_path}")
    except Exception as e:
        print(f"Warning: could not trim {png_path}: {e}")


def render_tikz_figures(text, base_dir):
    """
    Finds every \\begin{tikzpicture}...\\end{tikzpicture} block, compiles each
    to a PNG via a temporary article LaTeX file + pdflatex + pdftoppm + PIL trim,
    then replaces the block with \\includegraphics[width=\\columnwidth]{tikz_N.png}.

    Uses a large-page article document (no 'standalone' class required) and
    trims whitespace with PIL so the PNG is tight to the figure content.

    Must run AFTER strip_resizebox() so the tikzpicture is bare in the text.
    """
    tikz_pattern = re.compile(
        r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}',
        re.DOTALL
    )
    counter = [0]

    def replace_tikz(match):
        counter[0] += 1
        n = counter[0]
        tikz_code = match.group(0)

        standalone_tex = os.path.join(base_dir, f'tikz_{n}.tex')
        standalone_pdf = os.path.join(base_dir, f'tikz_{n}.pdf')
        standalone_png = os.path.join(base_dir, f'tikz_{n}.png')

        # Large-page article without standalone.cls; PIL crops whitespace afterwards.
        # lmodern + T1 + textcomp ensures >, <, >= render correctly in text nodes.
        tex_content = (
            r'\documentclass{article}' + '\n'
            r'\usepackage[T1]{fontenc}' + '\n'
            r'\usepackage{lmodern}' + '\n'
            r'\usepackage{textcomp}' + '\n'
            r'\usepackage[paperwidth=30cm,paperheight=30cm,margin=0.5cm]{geometry}' + '\n'
            r'\usepackage{tikz}' + '\n'
            r'\usetikzlibrary{shapes.geometric,arrows.meta,positioning}' + '\n'
            r'\pagestyle{empty}' + '\n'
            r'\begin{document}' + '\n'
            + tikz_code + '\n'
            + r'\end{document}' + '\n'
        )

        with open(standalone_tex, 'w', encoding='utf-8') as f:
            f.write(tex_content)

        try:
            subprocess.run(
                ['pdflatex', '-interaction=nonstopmode',
                 f'-output-directory={base_dir}', standalone_tex],
                check=True, capture_output=True, text=True
            )
            print(f"pdflatex compiled tikz_{n}.tex successfully")
        except subprocess.CalledProcessError as e:
            print(f"Warning: pdflatex failed for tikz_{n}.tex — leaving tikzpicture unchanged")
            print(e.stdout or '')
            return match.group(0)

        if not os.path.exists(standalone_pdf):
            print(f"Warning: {standalone_pdf} not produced — leaving tikzpicture unchanged")
            return match.group(0)

        convert_pdf_to_png(standalone_pdf, standalone_png, dpi=300)

        if not os.path.exists(standalone_png):
            print(f"Warning: {standalone_png} not produced — leaving tikzpicture unchanged")
            return match.group(0)

        # Trim the large white page down to just the figure
        _trim_png_whitespace(standalone_png)
        print(f"TikZ figure {n} rendered: {standalone_png}")

        # Clean up intermediate pdflatex artefacts
        for ext in ('.tex', '.pdf', '.aux', '.log'):
            p = os.path.join(base_dir, f'tikz_{n}{ext}')
            if os.path.exists(p):
                os.remove(p)

        return f'\\includegraphics[width=\\columnwidth]{{tikz_{n}.png}}'

    return tikz_pattern.sub(replace_tikz, text)


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def is_mono_run(run):
    """True if this run should use a monospace font (e.g. from \\texttt)."""
    try:
        if run.style and run.style.name:
            sname = run.style.name.lower()
            if any(kw in sname for kw in _MONO_STYLE_KEYWORDS):
                return True
    except Exception:
        pass
    fname = (run.font.name or '').lower()
    return bool(fname) and any(m in fname for m in _MONO_FONT_KEYWORDS)


def para_is_centered(para):
    """True if the paragraph has explicit center alignment in its XML pPr."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return False
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return False
    return jc.get(qn('w:val'), '').lower() == 'center'


def para_has_drawing(para):
    return bool(para._element.findall(f'.//{qn("w:drawing")}'))


def run_has_drawing(run):
    return bool(run._r.findall(f'.//{qn("w:drawing")}'))


def para_has_math(para):
    """True only for standalone display-math paragraphs (no body text alongside the equation).
    Inline math paragraphs ($...$) also contain oMath but have substantial text — those are
    body text and should be justified, not centered."""
    if not para._element.findall(f'.//{{{_MATH_NS}}}oMath'):
        return False
    # If the paragraph carries meaningful text (non-math runs), it is inline math in a
    # body paragraph — don't treat it as display math.
    return len(para.text.strip()) < 5


def add_page_number_to_header(section):
    """Inserts a right-aligned PAGE field into the section header."""
    header = section.header
    if not header.paragraphs:
        para = header.add_paragraph()
    else:
        para = header.paragraphs[0]

    # Clear everything except paragraph properties
    p_elem = para._element
    for child in list(p_elem):
        if child.tag != qn('w:pPr'):
            p_elem.remove(child)

    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)


def set_table_borders(table):
    """APA table borders: top/bottom rules on table, bottom rule on header row, no verticals."""
    tblPr = table._tbl.tblPr

    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    tblBorders = OxmlElement('w:tblBorders')

    for name, visible in [('top', True), ('bottom', True),
                           ('left', False), ('right', False),
                           ('insideH', False), ('insideV', False)]:
        b = OxmlElement(f'w:{name}')
        if visible:
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '8')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
        else:
            b.set(qn('w:val'), 'none')
        tblBorders.append(b)

    tblPr.append(tblBorders)

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)

            tcBorders = OxmlElement('w:tcBorders')

            bottom_b = OxmlElement('w:bottom')
            if r_idx == 0:
                bottom_b.set(qn('w:val'), 'single')
                bottom_b.set(qn('w:sz'), '8')
                bottom_b.set(qn('w:space'), '0')
                bottom_b.set(qn('w:color'), '000000')
            else:
                bottom_b.set(qn('w:val'), 'none')
            tcBorders.append(bottom_b)

            for bname in ['top', 'left', 'right']:
                b = OxmlElement(f'w:{bname}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)

            tcPr.append(tcBorders)


def _clear_para_runs(para):
    """Removes all content from a paragraph, leaving only the pPr element."""
    p_elem = para._element
    for child in list(p_elem):
        if child.tag != qn('w:pPr'):
            p_elem.remove(child)


def format_captions(doc):
    """Adds sequential Figure/Table numbers and applies APA 7th caption formatting.

    Pandoc strips LaTeX counter numbers from \\caption{} output, leaving only the
    title text in Image Caption / Table Caption style paragraphs.  This function
    walks the document in order, assigns Figure N / Table N labels, and reformats
    each caption as:  **Figure N**  (bold, 12 pt)  +  line-break  +  *title*  (italic).
    """
    fig_num = 0
    tbl_num = 0
    _cont_pat = re.compile(r'^Table\s+\d+\s+continued', re.IGNORECASE)

    for para in doc.paragraphs:
        style_name = para.style.name
        text = para.text.strip()

        if not text:
            continue

        if style_name == 'Image Caption':
            fig_num += 1
            label = f"Figure {fig_num}"
            caption_text = text

        elif style_name == 'Table Caption':
            if _cont_pat.match(text):
                # "Table N continued" — keep as-is, just normalise font
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 2.0
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            tbl_num += 1
            label = f"Table {tbl_num}"
            caption_text = text

        else:
            continue

        # Rebuild the paragraph: clear all runs, then add label + break + italic title
        _clear_para_runs(para)

        r_label = para.add_run(label)
        r_label.bold = True
        r_label.font.name = 'Times New Roman'
        r_label.font.size = Pt(12)
        r_label.font.color.rgb = RGBColor(0, 0, 0)

        r_break = para.add_run()
        r_break.add_break()
        r_break.font.name = 'Times New Roman'
        r_break.font.size = Pt(12)

        r_title = para.add_run(caption_text)
        r_title.italic = True
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(12)
        r_title.font.color.rgb = RGBColor(0, 0, 0)

        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.left_indent = Inches(0)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.keep_with_next = True


def insert_references_heading(doc):
    """Inserts a 'References' Heading 1 paragraph before the first Bibliography entry if absent."""
    if any(p.text.strip() == 'References' and p.style.name.startswith('Heading')
           for p in doc.paragraphs):
        return
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    for child in list(body):
        if child.tag == f'{{{ns}}}p':
            style_el = child.find(f'.//{{{ns}}}pStyle')
            if style_el is not None and style_el.get(f'{{{ns}}}val') == 'Bibliography':
                heading_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                pStyle = OxmlElement('w:pStyle')
                pStyle.set(qn('w:val'), 'Heading1')
                pPr.append(pStyle)
                heading_p.append(pPr)
                run_el = OxmlElement('w:r')
                t_el = OxmlElement('w:t')
                t_el.text = 'References'
                run_el.append(t_el)
                heading_p.append(run_el)
                body.insert(list(body).index(child), heading_p)
                print("Inserted missing 'References' Heading 1 before Bibliography entries.")
                return


def apply_docx_styling(docx_path, output_path):
    """Enforces APA 7th Edition layout, margins, typography, page headers, and table borders."""
    doc = Document(docx_path)

    # 1. Page size (A4) + margins (1 inch) + page-number header for every section
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        add_page_number_to_header(section)

    # 2. Base styles — set JUSTIFY as default for all body paragraph styles
    for sname in ('Normal', 'Body Text', 'First Paragraph', 'Body Text Indent'):
        try:
            s = doc.styles[sname]
            s.font.name = 'Times New Roman'
            s.font.size = Pt(12)
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            s.paragraph_format.line_spacing = 2.0
            s.paragraph_format.space_after = Pt(0)
            s.paragraph_format.space_before = Pt(0)
        except KeyError:
            pass

    _heading_configs = [
        ('Heading 1', WD_ALIGN_PARAGRAPH.CENTER, True,  False),
        ('Heading 2', WD_ALIGN_PARAGRAPH.LEFT,   True,  False),
        ('Heading 3', WD_ALIGN_PARAGRAPH.LEFT,   True,  True),
    ]
    for hname, align, bold, italic in _heading_configs:
        try:
            h = doc.styles[hname]
            h.font.name = 'Times New Roman'
            h.font.size = Pt(12)
            h.font.bold = bold
            h.font.italic = italic
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.alignment = align
            h.paragraph_format.line_spacing = 2.0
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
        except KeyError:
            pass

    # 3. Insert missing References heading before Bibliography entries
    insert_references_heading(doc)

    # 4. Paragraph pass
    # Track context flags for alignment/indent decisions
    in_references = False
    in_abstract = False
    in_title_page = True        # True until first Heading encountered
    body_title_done = False     # True after the APA-required body-page title is processed

    # The body-page title repeats the main title centered on the first page of body text.
    # Both title page (para 0) and body page title (para ~9) start with this prefix.
    TITLE_PREFIX = "Treatment Success and Failure"

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name

        # --- Skip Bibliography: handled in separate pass below ---
        if style_name == 'Bibliography':
            continue

        # --- Section heading detection ---
        if text == "References" and style_name.startswith("Heading"):
            in_references = True
            in_abstract = False
            in_title_page = False
            para.style = doc.styles['Heading 1']
            continue
        elif text == "Abstract" and (style_name.startswith("Heading") or style_name in ('Normal', 'Body Text')):
            in_abstract = True
            in_references = False
            in_title_page = False
            para.style = doc.styles['Heading 1']
            continue
        elif style_name.startswith("Heading"):
            in_abstract = False
            in_title_page = False
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_with_next = True
            for run in para.runs:
                if not run_has_drawing(run):
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # --- Figure / image paragraphs ---
        if para_has_drawing(para):
            para.paragraph_format.first_line_indent = Inches(0)
            para.paragraph_format.left_indent = Inches(0)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if not run_has_drawing(run) and run.text.strip():
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # --- Display math paragraphs ---
        if para_has_math(para):
            para.paragraph_format.first_line_indent = Inches(0)
            para.paragraph_format.left_indent = Inches(0)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            continue

        # --- Run-level font standardisation (monospace runs preserved) ---
        for run in para.runs:
            if run_has_drawing(run) or not run.text.strip():
                continue
            if is_mono_run(run):
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.name = 'Times New Roman'
                if run.font.size is None or run.font.size != Pt(12):
                    run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # --- Alignment and indentation by context ---
        if in_title_page:
            # Title page: all paragraphs centered, no indent
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Inches(0)
            para.paragraph_format.left_indent = Inches(0)

        elif in_abstract:
            if text.lower().startswith('keywords'):
                # Keywords line: label indented per APA; exit abstract after this paragraph
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0.5)
                in_abstract = False  # Body text (including the body-page title) follows next
            else:
                # Abstract body: justified, no paragraph indent
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0)

        elif in_references:
            # Reference entries (non-Bibliography style): hanging indent
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)

        elif style_name in ('Normal', 'Body Text', 'First Paragraph', 'Body Text Indent'):
            if not text:
                continue
            # Body-page title: APA repeats paper title centered at top of body text page
            if not body_title_done and text.startswith(TITLE_PREFIX):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.left_indent = Inches(0)
                body_title_done = True
            else:
                # Regular body text: justified, 0.5in first-line indent
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.left_indent = Inches(0)
        else:
            para.paragraph_format.first_line_indent = Inches(0)
            para.paragraph_format.left_indent = Inches(0)

    # 5. Bibliography pass — hanging indent + justify (separate because Pandoc omits the heading)
    for para in doc.paragraphs:
        if para.style.name == 'Bibliography':
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                if not run_has_drawing(run) and run.text.strip():
                    if not is_mono_run(run):
                        run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    # 6. Table borders and cell typography
    for table in doc.tables:
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.first_line_indent = Inches(0)
                    for run in para.runs:
                        if not run_has_drawing(run) and run.text.strip():
                            if not is_mono_run(run):
                                run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)

    # 7. APA caption formatting
    format_captions(doc)

    doc.save(output_path)
    print(f"Saved styled docx to: {output_path}")


def main():
    print("Step 1: Scan and convert all LaTeX PDF figures to high-resolution PNGs...")
    pdf_paths = []
    for root, dirs, files in os.walk(BASE_DIR):
        for file in files:
            if file.endswith(".pdf"):
                full_path = os.path.join(root, file)
                if file != "iCreate.pdf":
                    pdf_paths.append(full_path)

    for pdf_path in pdf_paths:
        png_path = os.path.splitext(pdf_path)[0] + ".png"
        convert_pdf_to_png(pdf_path, png_path, dpi=300)

    print("\nStep 2: [Skipped] Appendix B (iCreate.pdf) split is disabled...")

    print("\nStep 3: Inline LaTeX input files recursively...")
    full_latex_content = inline_latex(MAIN_TEX, BASE_DIR)

    print("\nStep 4: Preprocessing LaTeX markup for Pandoc...")

    # 4a. Replace .pdf figure inclusions with .png
    full_latex_content = re.sub(
        r'\\includegraphics(\[[^\]]*\])?\{([^}]+)\.pdf\}',
        r'\\includegraphics\1{\2.png}',
        full_latex_content
    )

    # 4b. Strip \resizebox wrappers (exposes bare tikzpicture for step 4d)
    print("Stripping \\resizebox from tables and figures...")
    full_latex_content = strip_resizebox(full_latex_content)

    # 4c. Strip LaTeX-only layout commands that produce Pandoc warnings/noise
    print("Stripping unsupported LaTeX layout commands...")
    full_latex_content = re.sub(r'\\FloatBarrier', '', full_latex_content)
    full_latex_content = re.sub(r'\\addtolength\{[^}]*\}\{[^}]*\}', '', full_latex_content)
    full_latex_content = re.sub(r'\\setlist\{[^}]*\}', '', full_latex_content)
    full_latex_content = re.sub(
        r'\\renewcommand\{\\'
        r'(?:topfraction|bottomfraction|textfraction|floatpagefraction|headrulewidth)'
        r'\}\{[^}]*\}',
        '', full_latex_content
    )
    full_latex_content = re.sub(
        r'\\setcounter\{(?:topnumber|bottomnumber|totalnumber)\}\{[^}]*\}',
        '', full_latex_content
    )

    # Strip \makeatletter...\makeatother blocks — these contain \@startsection
    # redefinitions that Pandoc cannot parse, causing all section headings to be
    # garbled with spacing-argument text like "1.5plus 0.2minus 0.1 0.5 Heading".
    full_latex_content = re.sub(
        r'\\makeatletter.*?\\makeatother', '', full_latex_content, flags=re.DOTALL
    )

    # Replace manual Abstract centering with a proper \section* so Pandoc emits
    # a Heading 1 paragraph (required for in_abstract detection to work).
    full_latex_content = re.sub(
        r'\\begin\{center\}\s*\\textbf\{Abstract\}\s*\\end\{center\}',
        r'\\section*{Abstract}',
        full_latex_content
    )

    # Strip \setlength{\parindent}{...} — layout command ignored by Pandoc
    full_latex_content = re.sub(r'\\setlength\{\\parindent\}\{[^}]*\}', '', full_latex_content)

    # 4d. Pre-render TikZ figures → PNG, replace with \includegraphics
    print("Pre-rendering TikZ figures to PNG...")
    full_latex_content = render_tikz_figures(full_latex_content, BASE_DIR)

    # 4e. Remove Appendix B entirely
    print("Removing Appendix B: iCreate Supplementary Material...")
    full_latex_content = re.sub(
        r'\\clearpage\s*\\section\{Appendix B: iCreate Supplementary Material\}'
        r'.*?\\includepdf\[[^\]]*\]\{iCreate\.pdf\}',
        '',
        full_latex_content,
        flags=re.DOTALL
    )

    with open(PREPROCESSED_TEX, 'w', encoding='utf-8') as f:
        f.write(full_latex_content)
    print(f"Created preprocessed LaTeX file: {PREPROCESSED_TEX}")

    print("\nStep 5: Running Pandoc to compile to draft DOCX...")
    pandoc_cmd = [
        "pandoc",
        PREPROCESSED_TEX,
        "--citeproc",
        f"--bibliography={BIB_FILE}",
        f"--csl={CSL_FILE}",
        "--wrap=none",
        "-o", DRAFT_DOCX,
    ]
    print(f"Running command: {' '.join(pandoc_cmd)}")
    try:
        subprocess.run(pandoc_cmd, check=True)
        print(f"Draft DOCX created: {DRAFT_DOCX}")
    except subprocess.CalledProcessError as e:
        print(f"Pandoc compilation failed: {e}")
        return

    print("\nStep 6: Enforcing APA 7th layout & styles via python-docx...")
    apply_docx_styling(DRAFT_DOCX, FINAL_DOCX)

    # Keep preprocessed LaTeX and draft DOCX for debugging if needed.
    print("\nStep 7: Done. Temporary files retained for debugging.")
    print("\nLaTeX to Lossless Word (DOCX) conversion complete!")
    print(f"Output file: {FINAL_DOCX}")


if __name__ == "__main__":
    main()
